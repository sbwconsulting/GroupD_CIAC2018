from copy import deepcopy
from decimal import Decimal
from decimal import InvalidOperation
from decimal import ROUND_HALF_UP
import logging
import re
import operator
from io import BytesIO
import os.path

from docx import Document
from docx.shared import Inches
from docx.oxml.shared import OxmlElement
from docx.oxml.shared import qn
from docx.enum.section import WD_ORIENTATION
#from docxtpl import DocxTemplate
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.utils.cell import get_column_letter
import numpy as np
import pandas as pd

from cpuc.utility import replace_strings_with_spaces
from cpuc.utility import swap_out_strings
from read_table_spec import gettableboundaries
from word_automation import autofit_table

DENSE_STYLE = 'dense' #must match style in excel design template
NORMAL_STYLE = 'normal'

class OldTable():
    """
    Deprecated. Original class used to generate tables 1, 2, and 4 in the
    D0 Preliminary Report.

    Top-level class for creation of CPUC report tables. Encapsulates
    behavior common to all tables.

    args:
    data      One or more pandas DataFrames from which to draw table data.
              These are stored internally as a list of dataframes, which is
              processed sequentially to fill all the rows needed in the table.
    colspec   dict specifying the data field used to generate the table's
              columns. Also contains the caption (display name) and value (value
              of the data field) for each column.

                  {
                      'field': 'db_field_name',
                      'columns': [
                          {
                              'caption': 'display name for column header',
                              'value': 'value in field by which to slice data',
                          },
                          ...
                      ]
                  }

              If a column has the caption name 'Total', no data is
              retrieved for that column. Instead, its value is the sum
              of the other values in its row.

    sections  dict specifying the sections of the table, and within each
              section, its rows.

                  {
                      'header': 'section header',
                      'rows': [
                          {
                              'index': 'display name for row index',
                              'fields: ['db_field1', ...],
                              'transforms': ['function', ...],
                              'reducer': function,
                          },
                          ...
                      ]
                  }

              Each row is represented by a row specification (row_spec),
              a dict describing the row's display name, the fields used to
              construct each value in the row, the transforms to be applied
              to derive each value, and a reducer function for combining
              multiple fields.

              Transforms are a list of functions in the Table class to
              apply to the data retrieved from each specified database
              field. Transform functions are called as a chain, passing
              the result of one transform to the next in the list. The
              last function should return a scalar value.

              The reducer function combines the transformed results for
              all the fields in the fields list into a single value. A
              reducer is not required if fields contains only a single
              database field.
    """
    def __init__(self, data, tablespec):
        self.dataframes = self._import_data(data)
        self.colspec = tablespec['colspec']
        self.sections = tablespec['sections']
        self.rows = []
        self.name = tablespec['name']
        self.title = tablespec['title']

    def _import_data(self, obj):
        dataframes = []
        if isinstance(obj, list):
            return [v for v in obj]
        else:
            return [obj]

    def to_list(self, formatted=False):
        self._build_table()
        return [row.to_list(formatted=formatted) for row in self.rows]

    def to_tuple(self, formatted=False):
        self._build_table()
        return tuple(row.to_tuple(formatted=formatted) for row in self.rows)

    def data_to_list(self):
        self._build_table()
        output = []
        for row in self.rows:
            # Skip empty rows.
            data = row.data_to_list()
            if any(data):
                output.append(data)
        return output

    def data_to_tuple(self):
        self._build_table()
        output = []
        for row in self.rows:
            # Skip empty rows.
            data = row.data_to_tuple()
            if any(data):
                output.append(data)
        return tuple(output)

    def _build_table(self, rebuild=False):
        """
        Build the rows that make up the table. If they've already been
        built, returns without doing anything.

        kwargs:
        rebuild  If True, rebuild the table from scratch. Otherwise, only
                 build the table if it hasn't already been built.
        """
        if len(self.rows) and not rebuild:
            return

        self.rows = []
        self.rows.append(
            OldRow(None, [c['caption'] for c in self.colspec['columns']])
        )
        for section in self.sections:
            self.rows.append(
                OldRow(section['header'], [None] * len(self.colspec['columns']))
            )
            for row in section['rows']:
                self.rows.append(
                    self._build_row(self.dataframes[0], row)
                )

    def _build_row(self, df, row_spec):
        """
        Given a dataframe and a row specification, return an OldRow
        object that represents the rendered contents of a table row.

        args:
            df        pandas dataframe
            row_spec  row specification dict
        """
        data = []
        total = 0
        for column in self.colspec['columns']:
            if column['caption'] == 'Total':
                # Assumption: Total is always the last column, so we can
                # accumulate it as we go...
                if len(row_spec['fields']) > 1:
                    # ...but for a percent change total, we can't simply add up
                    # the preceding values in the row. Rather, we need to calculate
                    # each value in the fields list, then apply the reducer to
                    # the results of those calculations.
                    values = self._calculatevalues(df, row_spec, column)
                    combinedvalue = self._apply_reducer(row_spec['reducer'], values)
                    data.append(combinedvalue)
                else:
                    data.append(total)
            else:
                values = self._calculatevalues(df, row_spec, column)

                if len(values) > 1:
                    combinedvalue = self._apply_reducer(row_spec['reducer'], values)
                else:
                    combinedvalue = values[0]
                total += combinedvalue
                data.append(combinedvalue)
        return OldRow(row_spec['index'], data)

    def _calculatevalues(self, df, row_spec, column):
        """
        Retrieve data for the fields that make up a row and column, apply
        transforms to them, and return a list of transformed scalar values.
        """
        values = []
        for field in row_spec['fields']:
            if column['value'] == 'total':
                column_data = df
            else:
                column_data = df[df[self.colspec['field']] == column['value']]
            value = self._getvalue(column_data, field)
            values.append(value)
        return values

    def _getvalue(self, df, field_spec):
        """
        Get the value of a particular data field from a dataframe, apply
        transforms, and return the result. Assumes that the transforms
        will produce a scalar value.
        """
        filtered = self._filter(df, field_spec)
        value = self._transform(
            filtered,
            field_spec,
        )
        if isinstance(value, (list, pd.DataFrame, pd.Series)):
            raise ValueError('Transforms must produce a scalar value')
        return value

    def _filter(self, data, field_spec):
        """
        Filters rows in a dataframe according to the filters described
        in a field spec. A filter method should return the same number
        of columns as its source dataframe.
        """
        try:
            filters = field_spec['filters']
        except KeyError:
            # Without any filters, default to retrieving the field spec's
            # column. Equivalent to specifying 'take_field' as the first
            # transform function.
            try:
                column = data[field_spec['field']]
            except KeyError:
                raise ValueError(
                    'Field %s does not exist in dataframe' % field_spec['field']
                )
            return column

        output = deepcopy(data)
        for filt in filters:
            output = getattr(self, filt)(output)
        return output

    def _transform(self, data, field_spec):
        """
        Applies transform methods to a dataframe, series, or scalar
        value. The result of a transform method should always be a
        scalar value.
        """
        output = deepcopy(data)
        for transform in field_spec['transforms']:
            output = getattr(self, transform)(output, field=field_spec['field'])
        return output

    def _apply_reducer(self, reducer, values):
        try:
            combinedvalue = reducer(values)
        except KeyError:
            raise ValueError(
                'A reducer function must be specified if the'
                ' row_spec contains multiple fields.'
            )
        return combinedvalue

    def sum_data(self, df, **kwargs):
        return df.sum()

    def is_sampled(self, df, **kwargs):
        #TODO Fix this. Doesn't account for Dropped or treatment or fuel
        #return df[df['SampledProject'].notna()]
        return df[df['SampleID'].notna()]

    def not_sampled(self, df, **kwargs):
        return df[df['SampledProject'].isna()]

    def all(self, df, **kwargs):
        return df

    def take_field(self, df, field=None, **kwargs):
        """
        Return a single field (column) from a dataframe.
        """
        return df[field]

    def project_count(self, df, **kwargs):
        """
        Return a count of the unique SBW_ProjID_Full values.
        """
        #TODO Fix this if it's needed
        #return len(df['SBW_ProjID_Full'].unique())
        return len(df.index)

    def count(self, df, **kwargs):
        """
        Return a count of the rows in a dataframe.
        """
        return len(df)

    def kWh_to_MWh(self, value, **kwargs):
        return value / 1000

    def kW_to_MW(self, value, **kwargs):
        return value / 1000

    @classmethod
    def percent_change(cls, values):
        """
        Return percentage change from first to second value.
        """
        msg = 'Table.percent_change requires exactly 2 values.'
        if not isinstance(values, (list, tuple)):
            raise ValueError(msg)
        if len(values) != 2:
            raise ValueError(msg)

        try:
            val = ((values[1] / values[0]) - 1) * 100
        except:
            return f'cannot do pct changes with {values[1]} and {values[0]}'
        else:
            return val
        # return ((values[1] / values[0]) - 1) * 100
        # return (1-(values[1] / values[0])) * 100

    @classmethod
    def percent_of_first(cls, values):
        """
        Return percentage of first value represented by second value.
        """
        msg = 'Table.percent_of_first requires exactly 2 values.'
        if not isinstance(values, (list, tuple)):
            raise ValueError(msg)
        if len(values) != 2:
            raise ValueError(msg)

        try:
            val = (values[1] / values[0]) * 100
        except:
            return f'cannot do pct of first with {values[1]} and {values[0]}'
        else:
            return val
        # return (values[1] / values[0]) * 100

    @classmethod
    def ratio(cls, values):
        """
        Return the ratio of first value to second.
        """
        msg = 'Table.ratio requires exactly 2 values.'
        if not isinstance(values, (list, tuple)):
            raise ValueError(msg)
        if len(values) != 2:
            raise ValueError(msg)

        try:
            val = values[0] / values[1]
        except:
            return f'cannot do ratio of {values[0]} and {values[1]}'
        else:
            return val
        # return values[0] / values[1]


class OldRow():
    """
    Deprecated. Original class used by OldTable to generate tables 1, 2,
    and 4 in the D0 Preliminary Report.

    Represents a table row's contents, including its index and values.
    """
    def __init__(self, index, data):
        self._index = index
        self._data = data

    @property
    def index(self):
        return self._index

    @property
    def data(self):
        return self._data

    @property
    def num_columns(self):
        return len(self._data + 1)

    def _convert_floats(self, remove_strings=False):
        """
        Convert numpy.float64 values to python float, so anything we
        pass the table to will be able to reformat values as needed. For
        example, String.format doesn't know what to do with a float64.

        This shouldn't lose any precision; python already uses a 64-bit
        float implementation.
        """
        data = []
        for value in self._data:
            if isinstance(value, np.float64):
                converted = float(value)
            elif remove_strings and isinstance(value, str):
                converted = None
            else:
                converted = value
            data.append(converted)
        return data

    def to_list(self, formatted=False):
        return [self._index] + self._convert_floats()

    def to_tuple(self, formatted=False):
        return (self._index, ) + tuple(self._convert_floats())

    def data_to_list(self):
        """
        Return only the numerical data cells in the row as a list.
        """
        return self._convert_floats(remove_strings=True)

    def data_to_tuple(self):
        """
        Return only the numerical data cells in the row as a tuple.
        """
        return tuple(self._convert_floats(remove_strings=True))


class Table():
    """
    Top-level class for creation of CPUC report tables. Encapsulates
    behavior common to all tables.
    """
    def __init__(self, **kwargs):
        self.type = None
        self.rows = []
        self.name = None
        self.num_rows = 0
        self.num_cols = 0
        self.col_widths = []
        self.comments = dict()
        self.datastart = dict()
        self.tablestart = dict()
        self.sheet = None    
        self.orientation = 'portrait'
        self.style = None
        self.autofit = None
        self._df = kwargs.get('df', None)

    @classmethod
    def from_spec(cls, spec, **kwargs):
        """
        Initialize a Table from a specification structure. The spec comes in as a TableSpec object       
        """
        table = Table(**kwargs)

        table.__dict__.update(spec.__dict__)

        if table.type == 'data':
            return table
        else:
            table.num_rows = spec.dimensions['rows']
            table.num_cols = spec.dimensions['columns']

            rows = deepcopy(spec.cells)
            for _, row in enumerate(rows):               
                new_row = []
                for col_index in range(1, table.num_cols + 1):
                    try:
                        if row[0][0]['col'] == col_index:
                            # Pop cell data off the row and create a new
                            # Cell from it.                            
                            new_row.append(Cell.from_list(row.pop(0), table=table))
                        else:
                            # Create empty cell
                            new_row.append(Cell.from_list(None))
                    except IndexError:
                        # No more cells in this row, pad the right with
                        # empty cells.
                        new_row.append(Cell.from_list(None))

                #TODO generalize this as this was specifically for D0
                # Condense first three columns into one with appropriate
                # heading formats
                '''
                keep_cell = new_row[0]                        
                for index in range(3):
                    cell = new_row.pop(0)
                    if cell.value:
                        keep_cell = cell
                        if index == 1:
                            # Second column gets indentation
                            keep_cell.style = 'tbl text 2'
                        elif index == 2:
                            # Third column gets bold and right alignment
                            keep_cell.style = 'tbl strong r'            
                new_row.insert(0, keep_cell)
                '''

                table.rows.append(new_row)                

        return table

    @classmethod
    def from_list(cls, rows, name, **kwargs):
        """
        Initialize a Table from a list of lists, where each item represents
        the contents of a cell in the table.

        args:
        rows  list of lists
        name  name of the table to use as its title

        kwargs:
        table  an existing Table object to build upon. Without specifying
               a table argument, from_list builds a new Table from scratch
        """
        table = kwargs.get('table', Table())
        if not table.name:
            table.name = name

        # Try not to destroy the original list
        rows_copy = deepcopy(rows)
        total_cols = 0
        for row in rows_copy:
            new_row = []
            for cell in row:
                new_row.append(Cell.from_list(cell))
            table.rows.append(new_row)
            num_cols = len(new_row)
            if num_cols > total_cols:
                total_cols = num_cols
        if not table.num_cols:
            table.num_cols = total_cols
        if not table.num_rows:
            table.num_rows = len(table.rows)
        return table

    @property
    def df(self):
        if self._df is None:
            logging.warning("Empty dataframe for Table '%s'" % self.name)
        return self._df

    def to_list(self):
        return [[cell.str_value for cell in row] for row in self.rows]

    def to_data_list(self, rstart=0, cstart=0):        
        tmp = [[cell.value for cell in row if cell.col is not None and cell.row is not None and cell.col >=cstart and cell.row >=rstart] for row in self.rows]
        return tmp[rstart:]

    def to_numeric_list(self):
        return [[cell.value for cell in row] for row in self.rows]

    def to_dict(self):
        """
        Write out the table as a dict.
        """
        pass

    def to_word(self, sfsession, template, outfile):
        """
        Write the table to a Word file.

        args:
        template  Word file to open as a template. Required for custom styles
                  to work correctly.

        outfile   Output file to write to.
        """        
        doc = Document(template)

        #set orientation
        if self.orientation == 'landscape':
            section = doc.sections[-1]
            new_width, new_height = section.page_height, section.page_width
            section.orientation = WD_ORIENTATION.LANDSCAPE #or 1
            section.page_width = new_width
            section.page_height = new_height

        
        #add caption
        # doc.add_paragraph(style='tbl pre')        
        doc.add_paragraph(self.name, style='Caption table' )
        # Add table.
        word_table = doc.add_table(self.num_rows, self.num_cols)

        # Add border styling.
        self.set_table_border(
            word_table,
            top={'sz': 12, 'val': 'single', 'color': '#789b68'},
            insideH={'sz': 4, 'val': 'single', 'color': '#a5a5a5'},
            bottom={'sz': 12, 'val': 'single', 'color': '#789b68'},
        )

        # Zero out cell margins.
        self.set_table_cell_margin(word_table, 0)

        # Fill table cells.
        for row_index, row in enumerate(self.rows):
            headerrow = False
            for col_index, cell in enumerate(row):
                if not cell.col: #To skip merged cells
                    continue
                if not cell.row:
                    continue
                word_cell = word_table.cell(row_index, col_index)                
                #to clear out non=breaking spaces for wrapping cells
                # if not cell.wraptext:
                word_cell.text = cell.str_value
                # else:
                #     word_cell.text = ' '.join(cell.str_value.split())
                if cell.commentid is not None:
                    commentid = word_cell.paragraphs[0].add_run(str(cell.commentid))
                    commentid.font.superscript = True
                # word_cell.paragraphs[0].style = cell.style
                word_cell.paragraphs[0].style = convert_table_style(cell.style, self.style)
                # Apply background shading. 
                if cell.background:
                    self.shade_cells([word_cell], cell.background)               
                if cell.style and 'head' in cell.style:
                    word_cell.vertical_alignment = 1 #wdCellAlignVerticalCenter
                    headerrow = True

                # Merge colspan cells
                if cell.mergespan:
                    word_cell.merge(
                        word_table.cell(
                            row_index + int(cell.rowspan) - 1,
                            col_index + int(cell.colspan) - 1,
                        )
                    )
                    if cell.rowspan > 1:
                        rowspan = cell.rowspan


                # Adjust column widths. Must be done here rather than at
                # column level, because Word ignores column level.
                # https://stackoverflow.com/questions/43051462/python-docx-how-to-set-cell-width-in-tables
                # if self.col_widths:
                if cell.colwidth:
                    # word_cell.width = Inches(self.col_widths[col_index])
                    #this doesn't seem to have any effect on the final doc.
                    word_cell.width = Inches(float(cell.colwidth))
                    self.autofit = 'window'
            #set row to repeat if header
            if headerrow:
                self.set_repeat_table_header(word_table.rows[row_index])
        # self.set_table_autofit(word_table, 'content') #this 
        # if self.autofit != 'content':
        # self.set_table_autofit(word_table, self.autofit)
        # word_table.autofit = False
        
        # #not sure why this is here so removing it for now
        # tblW = word_table._tbl.tblPr.first_child_found_in('w:tblW')
        # tblW.set(qn('w:type'), 'dxa')
        # tblW.set(qn('w:w'), '2340')
        
        if (self.comments):
            # if len(self.comments)>1:
            # sorted_comments = sorted(self.comments['comments'].items(), key=operator.itemgetter(1))
            sorted_comments = sorted(self.comments.items(), key=operator.itemgetter(1))
            # else:
                # sorted_comments = self.comments
            for c in sorted_comments:
                p = doc.add_paragraph(f'{c[1]}\t{c[0]}')
                p.style='tbl note'
            if 'p' in locals():
                p.style = 'tbl note last'
        else:
            doc.add_paragraph(style='tbl post')

        # Remove empty first paragraph.
        # https://github.com/python-openxml/python-docx/issues/33
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)
        p._p = p._element = None

        # if isinstance(outfile, str):
        file_io = BytesIO()
        doc.save(file_io)
        try:            
            doc.save(file_io)
        except PermissionError:
            # if isinstance(outfile, str) and 'permission denied' in e.strerror.lower():
            # if isinstance(outfile, str) and 'permission denied' in e.args[1].lower():
            tmpoutfile = outfile.replace('.doc', 'tmp.doc')
            msg = f'bone head! Close {outfile} nexttime. I save you this time. Look at {tmpoutfile}'
            logging.warning(msg)
            print(msg)
            doc.save(tmpoutfile)
        except Exception as e:
                msg = (f'problem saving {outfile}: {e}')
                logging.critical(msg)
                print(msg)

        filename = os.path.basename(outfile) #.replace('.csv', '_raw.csv')
        pathroot = outfile.split(filename)[0]
        folder_item = sfsession.get_item_by_local_favorites_path(pathroot)
        sfsession.upload_file(folder_item.id, filename, file_io)


    def to_excel(self, wb):
        """ 
        Write the table to an excel spreadsheet file.

        args:
        wb  The workbook to write to. Does not save the file, only writes to it.
        """

        
        ws = wb[self.sheet]
        
        if self.type == 'data':
            write_data(ws, self)    
        else:
            data = self.to_data_list(self.datastart['row']-self.tablestart['row']-1,self.datastart['col']-self.tablestart['col'])
            #print(data)
            #print()
            write_worksheet(ws, data, rStart=self.datastart['row'], cStart=self.datastart['col'])
            hide_spec(ws, rStart=self.tablestart['row'], cStart=self.tablestart['col'])

    def post_process(self, wrdapp, wordfile):
        """
        Post process the word files to fix things openpyxl couldn't do :(
        """
        operations = self.autofit
        result = autofit_table(wrdapp, wordfile, operations)
        if not result:
            msg = f'Problem post processing {wordfile}'
            print(msg)
            logging.critical(msg)

    def set_table_cell_margin(self, table, margin):
        """
        Set the default cell margin for a table.

        args:
        table   python-docx Table object
        margin  size of margin in twentieths of a point (1/1440 of an inch)
        """
        tblPr = table._tbl.tblPr
        tblCellMar = tblPr.first_child_found_in('w:tblCellMar')
        if tblCellMar is None:
            tblCellMar = OxmlElement('w:tblCellMar')
            tblPr.append(tblCellMar)

        for side in ('start', 'top', 'end', 'bottom'):
            element = OxmlElement('w:{}'.format(side))
            element.set(qn('w:w'), str(margin))
            tblCellMar.append(element)

    def shade_cells(self, cells, shade):
        """
        Apply background shading to a Word table cell.
        https://github.com/python-openxml/python-docx/issues/146

        args:
        cells  list of python-docx cells
        shade  hex color code as a string (for example, '#ff0000')
        """
        for cell in cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcVAlign = OxmlElement('w:shd')
            tcVAlign.set(qn('w:fill'), shade)
            tcPr.append(tcVAlign)

    def set_table_border(self, table, **kwargs):
        """
        Apply border styling to a Word table.
        https://stackoverflow.com/questions/33069697/how-to-setup-cell-borders-with-python-docx

        args:
        table   python-docx Table object

        kwargs:
        start   style dict for left outer border
        top     style dict for top outer border
        bottom  style dict for bottom outer border
        end     style dict for right outer border
        insideH style dict for horizontal inner borders
        insideV style dict for vertical inner borders

        Each style dict may contain the following values:
        sz      width of the border in eighths of a point
        val     style of the border. 'single' is a single line.
                For other possible values, see
                http://officeopenxml.com/WPtableBorders.php
        color   hex color code as a string (for example, '#ff0000')
        space   spacing offset in points
        shadow  if 'true', add a shadow effect to the border
        """
        tblPr = table._tbl.tblPr
        tcBorders = tblPr.first_child_found_in('w:tblBorders')
        if tcBorders is None:
            tcBorders = OxmlElement('w:tblBorders')
            tblPr.append(tcBorders)

        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                for key in ('sz', 'val', 'color', 'space', 'shadow'):
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

    def set_repeat_table_header(self, tablerow):
        """ 
        from https://github.com/python-openxml/python-docx/issues/322
        
        set repeat table row on every new page
        """
        tr = tablerow._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)
        # return tablerow
    
    def set_table_autofit(self, table, fitsetting='content'):
        """
        Set a table to be either fit to contents or fit to window
        The fit setting should be either fixed, content, or window
        """

        tblPr = table._tbl.tblPr
        tblfit = OxmlElement('w:tblW')
        if fitsetting == 'window':
            # <w:tblW w:w="5000" w:type="pct"/>
            tblfit.set(qn('w:type'), 'pct')
            tblfit.set(qn('w:w'), '5000') #5000 is the full width
            tblPr.append(tblfit)
        elif fitsetting == 'fixed':
            #don't know setting for fixed yet. so setting using docx built in
            table.allow_autofit = False
        else: # fitsetting == 'content':
            # <w:tblW w:w="0" w:type="auto"/>
            # tblfit.set(qn('w:type'), 'auto')
            # tblfit.set(qn('w:w'), '0')
            # tblPr.append(tblfit)
            table.allow_autofit = True

class Cell():
    """
    Represents an individual cell in a table, including its properties
    and value. Given the right metadata, can calculate its value based
    on a pandas DataFrame.

    The simplest Cell is empty:

        cell = Cell()
        cell.value  # None

    A text-only cell used for captions:

        cell = Cell(value='Claim')
        cell.value  # 'Claim'

    Formatting may be applied:

        cell = Cell(
            value='% change from Claim',
            align='right',
            bold=True,
            background='#88FF00',
        )

    Specify a colspan of `n` to span across the next `n` cells.
    Alignment is applied across the spanned columns:

        cell = Cell(
            value='Claimed Lifetime Net Savings',
            colspan=4,
            align='center',
        )

    NOT IMPLEMENTED: Build the cell from an openpyxl worksheet and cell
    reference:

        wb = load_workbook('workbook.xlsx')
        cell = Cell.from_excel(wb['Sheet1'], 'C6')
    """
    CALC_REGEX = 'calc\((.+)\)'

    def __init__(self, **kwargs):
        self.excel = kwargs.get('excel', None)
        self._value = kwargs.get('value', None)
        #self.source = kwargs.get('source', None)
        self.alignment = kwargs.get('align', None)
        self.background = kwargs.get('background', None)
        self.mergespan = kwargs.get('mergespan', None)
        self.calculation = kwargs.get('calculation', None)
        self.style = kwargs.get('style', 'tbl text')
        self.format = kwargs.get('format', None)
        self.wraptext = kwargs.get('wraptext', None)
        # self.style = kwargs.get('style', None)
        self.decimals = kwargs.get('decimals', 0)
        self.percent = kwargs.get('percent', False)
        self.row = kwargs.get('row', None)
        self.col = kwargs.get('col', None)
        self.fields = kwargs.get('fields', [])
        self.filters = kwargs.get('filters', [])
        self.colfields = kwargs.get('colfields', [])
        self.colfilters = kwargs.get('colfilters', [])
        self.transforms = kwargs.get('transforms', [])
        self.reducer = kwargs.get('reducer', None)
        self.commentid = kwargs.get('commentid', None)
        self.colwidth = kwargs.get('colwidth', None)

        # Reference back to the containing table
        self.table = kwargs.get('table', None)
        table_df = None
        if self.table is not None and self.table.df is not None:
            table_df = self.table.df
        self._df = kwargs.get('df', table_df)

        if self.mergespan:
            # if self.mergespan['columns']>1:
            self.colspan = self.mergespan['columns']
            # if self.mergespan['rows']>1:
            self.rowspan = self.mergespan['rows']
            

    @classmethod
    def from_list(cls, data, **kwargs):
        """
        Initializes a cell from a list that contains a metadata dict
        followed by a string or number that represents the cell's data.
        Instead of a list, may also be a single scalar value, which is
        used for the cell's data without any attached metadata. An empty
        cell is represented by the value None.

        Numbers are stored as Decimal so they can later be accurately
        rounded.

        Example valid data inputs:
            [{'style': 'tbl head', 'colspan': 3}, 'Cumulative Thingies']

            [{'bold': True}, 5.2929]

            42

            None

        kwargs:
        table  a Table object that this Cell should be part of
        """
        if not data:
            return Cell(value=None)
        try:
            metadata = deepcopy(data[0])
        except TypeError:
            return Cell(value=data, table=kwargs.get('table'))
        else:
            value = data[1]

        # Override with passthrough value.
        text = metadata.get('text')
        passthru = False
        if text is not None:
            value = text
            ctype = metadata.get('type')
            if ctype == 'di':
                passthru = True

        # Set default cell style. Numbers get right alignment, strings
        # get left.                
        try:
            as_number = Decimal(value)
        except InvalidOperation:
            as_number = None
        except TypeError:
            if metadata.get('fields', None):
                as_number = True
            else:                
                as_number = False

        if not metadata.get('style'):
            if as_number:
                metadata['style'] = 'tbl text r'
            else:
                metadata['style'] = 'tbl text'
        if as_number and passthru:
            #set decimals here
            decimals = metadata.get('decimals')
            if decimals is not None:
                if decimals == 0:
                    value = str(int(round(float(value))))
                else:
                    value = str(round(float(value), int(decimals)))

        # Break properties into lists.
        Cell._split_dict(metadata, 'fields')
        for prop in ['filters', 'transforms']:
            Cell._split_dict_twice(metadata, prop)

        return Cell(value=value, table=kwargs.get('table'), **metadata)

    @classmethod
    def _split_dict(cls, dikt, key):
        """
        Retrieve a value from a dictionary, split it into a list based
        on comma delimiter, and replace the dictionary value with the
        resulting list.
        """
        value = dikt.get(key)
        if value is not None:
            dikt[key] = value.split(', ')

    @classmethod
    def _split_dict_twice(cls, dikt, key):
        """
        Retrieve a value from a dictionary, split it into a list based
        on comma delimiter. Then break those list items into lists based
        on the delimiter ' and '. Replace the dictionary value with the
        list of lists.
        """
        value = dikt.get(key)        
        if value is not None:
            #replace any quoted text that may have internal bits that would mess this up        
            repl = replace_strings_with_spaces(value)
            replacements = repl['replace_dict']
            value = repl['newstring']
            comma_list = value.split(', ')                    
            and_list = list(map(lambda x: x.split(' and '), comma_list))
            modand_list = []
            for item in and_list:
                item = swap_out_strings(item, replacements)
                modand_list.append(item)
            dikt[key] = modand_list

    @classmethod
    def parse(cls, value):
        """
        Initializes a cell from a value, parsing it for formatting and
        calculations. Numbers are stored as Decimal so they can later
        be accurately rounded.
        """
        if not value:
            return Cell(value=None)

        text = str(value)
        # styles = []
        if re.search(r'\|', text):
            format_text, content = text.split('|')
        else:
            format_text, content = '', text

        format_list = format_text.split(', ')
        formats = {}
        for fmt in format_list:
            if re.search(': ', fmt):
                directive, value = fmt.split(': ')
                try:
                    value = Decimal(value)
                except InvalidOperation:
                    value = str(value)

                if value == 'True':
                    value = True
                if value == 'False':
                    value = False

                formats[directive] = value

        # Set default cell style. Numbers get right alignment, strings
        # get left.
        try:
            as_number = Decimal(content)
        except InvalidOperation:
            as_number = None

        if not formats.get('style'):
            if as_number:
                formats['style'] = 'tbl text r'
                content = as_number
            else:
                formats['style'] = 'tbl text'

        return Cell(value=content, **formats)

    @property
    def df(self):
        if self._df is None:
            logging.warning('Empty dataframe for Cell (%d, %d)' % (self.row, self.col))
        return self._df

    @property
    def value(self):
        """
        Return the value of the cell. An explicit value always overrides
        a calculated value.
        """
        if self._value is not None:
            return self._value
        if self.calculation is not None:
            return self._calculate()
        if self.fields:
            value = self._calculate_from_fields()
            try:
                strvalue = str(value)
            except:
                #  strvalue = None
                 strvalue = 'NA'
            if strvalue.lower() == "nan":
                value = "NA" 
            return value
        else:
            return None

    @property
    def str_value(self):
        """
        Return the value of the cell as a string. Returns an empty
        string if the cell is blank. For numeric values, rounds to
        the correct number of digits after the decimal, according to
        the cell's `decimals` property. Also adds commas as thousands
        separators and applies any percent formatting.
        """
        value = self.value
        if value is not None:
            if self.format !='@':
                if isinstance(value, np.integer):                
                    value = int(value)            
                try:
                    quantizer = Decimal('0.{}'.format('0' * self.decimals))
                    value = Decimal(value).quantize(
                        quantizer,
                        ROUND_HALF_UP,
                    )
                except InvalidOperation:
                    # value = value.replace(' ', '\u00a0') #add in nonbreaking space
                    pass
                else:
                    # 2011 is the hard hyphen to keep negative values from separating from values in narrow columns. 
                    # Jeff J says it causes it to get stuck in Times New roman, but he will live with that.
                    # if '-' in str(value):
                    #     print('hypehn')
                    value = '{:,.{}f}'.format(value, self.decimals).replace('-', '\u2011') #.replace('-', '\u2212')u'\u2014' '\u2013'

                if self.percent:
                    value = '{}%'.format(value)

            return str(value)
        else:
            return ''

    def _calculate(self):
        """
        Calculate the value of the cell based on a calculation
        expression.
        """
        calc = Calculation(self.calculation, df=self.df)
        try:
            self._value = calc.value
        except RuntimeError:
            self._value = '#ERROR'

        return self._value

    def _calculate_from_fields(self):
        """
        Calculate the value of the cell based on fields, filters,
        transforms, and reducers.
        """
        # Nothing to calculate with
        if self.df is None:
            return None

        values = self._collect_values()
        if len(values) > 1:
            value = self._apply_reducer(values, self.reducer)
        else:
            value = values[0]
        self._value = value
        return self._value

    def _collect_values(self):
        """
        Collect data for the fields that make up the cell's data. Filter
        values to a single column if necessary.
        """
        df = self.df

        # For totals, exclude column filtering to get the full count
        # from the dataframe. For example, each PA column in Table 1
        # should be filtered to that particular PA, but the Total column
        # should reflect counts across all PAs.
        if self.colfilters == 'total': #hmm this may need to change
            column_data = df
        elif self.colfilters is not None:
            column_data = df[df[self.colfields] == self.colfilters]
        else:
            column_data = df

        values = []
        for index, field in enumerate(self.fields):
            if self.filters is not None:
                if len(self.fields) == len(self.filters):
                    try:
                        filters = self.filters[index]
                    except IndexError:
                        raise ValueError(
                            'Field %s is missing a corresponding filter' % field
                        )
                else:
                    filters = self.filters[0] * len(self.fields)
            else:
                filters = None
            if self.transforms is not None:
                try:
                    transforms = self.transforms[index]
                except IndexError:
                    raise ValueError(
                        'Field %s is missing corresponding transforms' % field
                    )
            else:
                transforms = None

            value = self._getvalue(column_data, field, filters, transforms)
            values.append(value)
        return values

    def _getvalue(self, data, field, filters, transforms):
        """
        Filter a dataframe, extract a field (column) from it, apply
        transforms, and return the result. Assumes that the transforms
        will produce a scalar value.
        """
        filtered = self._filter(data, filters)
        if field == 'ClaimID':
            field_data = filtered
        else:
            if field in filtered:
                field_data = filtered[field]
            else:
                msg = f'field {field} must be in the data'
                logging.critical(msg)
                print(msg)
                # raise ValueError(f'field {field} must be in the data')
                return f'missing {field}'
        transformed = self._transform(field_data, transforms)
        if isinstance(transformed, (list, pd.DataFrame, pd.Series)):            
            raise ValueError('Transforms must produce a scalar value')
        return transformed

    def _filter(self, data, filters):
        """
        Filters rows in a dataframe according to a filter value from a
        Cell's properties. A filter method should return the same number
        of columns as its source dataframe.
        """
        result = data
        try:
            for filt in filters:
                #string replace in case of spaces in quoted items
                repl = replace_strings_with_spaces(filt)
                replacements = repl['replace_dict']
                filt = repl['newstring']
                if filt != 'None':
                    if len(filt.split(" ",2))==3:
                        filter_field, operator, value = filt.split(" ",2)
                        #put original string back
                        filter_field = swap_out_strings(filter_field, replacements)
                        value = swap_out_strings(value, replacements)
                        query = "result[result['{}'] {} {}]".format(
                            filter_field,
                            operator,
                            value,
                        )
                    elif len(filt.split())==2: #for things like .isnull()
                        filter_field, operator = filt.split()
                        #put original string back
                        filter_field = swap_out_strings(filter_field, replacements)
                        value = swap_out_strings(value, replacements)
                        query = "result[result['{}']{}]".format(
                            filter_field,                            
                            operator,
                        )
                    else:
                        msg = 'Filter is malformed'
                        raise ValueError(msg)
                    try:
                        result = eval(query, None, locals())
                    except Exception as e:
                        msg = f'filter query problem {query}, {e}'
                        logging.critical(msg)
                        print(msg)
                        return msg
        except TypeError:
            # filters is None, so return data untouched
            pass
        return result

    def _transform(self, data, transforms):
        """
        Applies transform methods to a dataframe, series, or scalar
        value. The result of a transform method should always be a
        scalar value.
        """
        output = data.copy()
        if transforms is None:
            #A standin to do nothing
            transforms = ['sum']
        for transform in transforms:
            func_name = 'x_{}'.format(transform)
            output = getattr(self, func_name)(output)
        return output

    def _apply_reducer(self, values, reducer):
        func_name = 'r_{}'.format(reducer)
        try:
            return getattr(self, func_name)(values)
        except AttributeError:
            raise ValueError(
                'A reducer function must be specified if the'
                ' cell contains multiple fields.'
            )

    """
    Transform functions
    """
    def x_count(self, df):
        """
        Return a count of the rows in a dataframe.
        """
        return len(df)

    def x_kW_to_MW(self, value):
        return value / 1000

    def x_kWh_to_MWh(self, value):
        return value / 1000

    def x_mmbtu(self, value):
        return value * 0.1

    def x_to_pct(self, value):
        return value * 100

    def x_project_count(self, df):
        """
        Return a count of the unique SBW_ProjID_Full values.
        """
        return len(df['SBW_ProjID_Full'].unique())

    def x_sum(self, df):
        """
        Return sum of values in dataframe.
        """
        return df.sum()

    def x_max(self, df):
        """
        Return max of values in dataframe.
        """
        return df.max()

    def x_nothing(self, value):
        """
        Return max of values in dataframe.
        """
        return value


    """
    Reducer functions
    """
    def r_percent_change(self, values):
        """
        Return percentage change from first to second value.
        Uses (1-(value2/value1))*100
        """
        msg = 'Table.percent_change requires exactly 2 values.'
        if not isinstance(values, (list, tuple)):
            raise ValueError(msg)
        if len(values) != 2:
            raise ValueError(msg)
        if values[0] == 0:  #to prevent divide by zero error
            return 'NA'     #may want to set this to 0
        try:
            val = ((values[1] / values[0]) - 1) * 100
        except:
            return f'Cannot calculate pct change with {values[0]} and {values[1]}'
        else:
            return val
        # return ((values[1] / values[0]) - 1) * 100
        # return (1-(values[1] / values[0])) * 100

    def r_percent_of_first(self, values):
        """
        Return percentage of first value represented by second value.
        """
        msg = 'percent_of_first reducer requires exactly 2 values.'
        if not isinstance(values, (list, tuple)):
            raise ValueError(msg)
        if len(values) != 2:
            raise ValueError(msg)
        if values[0] == 0:  #to prevent divide by zero error
            return None     #may want to set this to 0        
        
        try:
            val = (values[1] / values[0]) * 100
        except:
            return f'cannot do pct of first with {values[1]} and {values[0]}'
        else:
            return val
        # return (values[1] / values[0]) * 100

    def r_difference(self, values):
        """
        Return the difference of the two values, i.e. v1 - v2
        """
        msg = 'difference reducer requires exactly 2 values.'
        if not isinstance(values, (list, tuple)):
            raise ValueError(msg)
        if len(values) != 2:
            raise ValueError(msg)
        return (values[1] - values[0])

    def r_sum(self, values):
        """
        Return the sum of the two values, i.e. v1 - v2
        """
        msg = 'sum reducer requires exactly 2 values.'
        if not isinstance(values, (list, tuple)):
            raise ValueError(msg)
        if len(values) != 2:
            raise ValueError(msg)
        return (values[1] + values[0])

    def r_ratio(self, values):
        """
        Return the ratio of first value to second.
        """
        msg = 'ratio reducer requires exactly 2 values.'
        if not isinstance(values, (list, tuple)):
            raise ValueError(msg)
        if len(values) != 2:
            raise ValueError(msg)
        if values[1] == 0:  #to prevent divide by zero error
            return None     #may want to set this to 0 or something else

        try:
            val = values[0] / values[1]
        except:
            return f'Cannot do ratio with {values[0]} and {values[1]}'
        else:
            return val
        # return values[0] / values[1]

    @classmethod
    #def from_excel(cls, worksheet, cell):
    def from_excel(self, worksheet, cell):
        """
        Return a Cell based on data from an Excel worksheet and cell.
        """
        # TODO: retrieve merged cells
        #       if cell is the start of a merged range, give it colspan
        #       if cell is elsewhere in merged range, give it None value
        #       apply alignment, text formatting, background
        #       if cell contents are calculation, fill calculation property
        #           else fill value property with cell value
        pass


class Calculation():
    """
    Handles parsing of a calculation from input text, preparation of
    calculation for execution, execution of the calculation, and return
    of the calculation's value.

    A calculation looks like a function call to a function named `calc`:

        calc(len('a'))

    Within a calculation string, {df} is replaced by the pandas dataframe
    with which the calculation was initialized:

        calc({df}['col'].sum())

    The previous example would return the sum of the dataframe's `col`
    column.

    args:
    df    pandas dataframe to use in the calculation

    text  calculation string to be executed
    """
    def __init__(self, text, **kwargs):
        self.text = text
        self._df = kwargs.get('df', None)

    @property
    def df(self):
        if self._df is None:
            logging.warning("Empty dataframe for calculation '%s'" % self.text)
        return self._df

    def _parse(self, text):
        """
        Parses the calculation from a calc(...) string.
        """
        try:
            match = re.match(r'calc\((.+)\)', text).group(1)
        except AttributeError:
            msg = '`calc()` is empty or missing from calculation string %s'
            logging.warning(msg % text)
            raise ValueError(msg % text)
        return match

    def _prepare(self, text):
        """
        Adds necessary variable context so a calculation's result
        can be extracted.
        """
        return text.format(df='self.df')

    def _execute(self, code):
        """
        Executes a calculation and returns its value.
        """
        try:
            result = eval(code, None, locals())
        except:
            msg = 'Calculation failed for %s'
            logging.warning(msg % code)
            raise ValueError(msg % code)
        return result

    @property
    def value(self):
        """
        Returns the calculated value.
        """
        return self._execute(self._prepare(self._parse(self.text)))

def write_worksheet(ws, data, rStart=1, cStart=1):
    """
    Write a list of lists (or tuple of tuples) to cells in a worksheet.
    """
    for rowindex, row in enumerate(data):
        for colindex, col in enumerate(row):
            try:
                ws.cell(row=(rowindex + rStart), column=(colindex + cStart)).value = col
            except Exception as e:
                #know this happens for merged cells (except first cell of merged group)
                print(f'error putting {col} in cell {rowindex + rStart}, {colindex + cStart}. {e}') 
                pass

def hide_spec(ws, rStart=1, cStart=1):
    """
    Hide the rows, cols around the table with the specification information
    """
    #get table bounds
    tablebounds = gettableboundaries(ws)
    #hide above and left of table
    if not tablebounds:
        return
    try:
        ws.row_dimensions[tablebounds['rows'][0]].hidden= True
        ws.column_dimensions[get_column_letter(tablebounds['cols'][0])].hidden= True
    except:
        pass
    #hide rows below table
    ws.row_dimensions.group(start=tablebounds['rows'][1] + 1, end=ws.max_row, hidden=True)
    
    #hide columns to the right of the table
    ws.column_dimensions.group(start=get_column_letter(tablebounds['cols'][1] +1), end=get_column_letter(ws.max_column), hidden=True)

def write_data(ws, tbl):
    """
    write a the fields of the dataframe to the worksheet
    """
    fields = tbl.fields
    rstart = tbl.tablestart['row']
    cstart = tbl.tablestart['col']
    roffset = 0
    if not fields:
        df = tbl.df
    else:
        try:
            df = tbl.df[fields]
        except Exception as e:
            print(f'{e}')
            logging.critical(f'for table {tbl.name} {e}')
            okfields = [x for x in fields if x in tbl.df.columns]
            try:
                df = tbl.df[okfields]
            except:
                return False
    for r in dataframe_to_rows(df):
        if type(r) is not list: #to skip the frozen list element between the header and data portion
            continue
        coffset = 0
        for value in r[1:]: #first column is always a blank column due to how df_to_rows works
            ws.cell(rstart + roffset, cstart + coffset).value = value
            coffset += 1
        roffset += 1
    
def convert_table_style(stylestring, style):
    """
    set the style to the passed style (dense or normal as defined in the SBW template)
    returns the updated style sting
    """

    text = stylestring
    dm = ' sm'

    if style == NORMAL_STYLE:
        text = text.replace(dm, '')
        return text

    if style == DENSE_STYLE:
        if dm in text:
            return text
        else:
            if ' ' == text[-2]:
                text = text[:-2] + dm + text[-2:]
            else:
                text = text + dm
            return text

    return text
